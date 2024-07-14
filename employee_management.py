import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document
from docx.shared import Inches
import pandas as pd
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import os
import csv
from datetime import datetime
import pytesseract
from PIL import Image
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

class User:
    def __init__(self, email, password):
        self.email = email
        self.password = password

class MonthlyReportSystem:
    def __init__(self, root):
        self.root = root
        self.root.title("Monthly Report Of SST")
        self.root.geometry("1200x700")
        self.root.configure(bg="#0F3057")
        
        self.style = ttk.Style()
        self.style.configure('TLabel', background='#DAE1E7', font=('Helvetica', 12))
        self.style.configure('TButton', font=('Helvetica', 12))
        self.style.configure('Header.TLabel', font=('Helvetica', 20, 'bold'))
        
        # Title
        title_label = ttk.Label(self.root, text="Monthly Report Of SST", style='Header.TLabel', anchor="center", background='#0F3057', foreground='#FFFFFF')
        title_label.pack(side=tk.TOP, fill=tk.X, pady=20)
 
        # Welcome Label
        self.welcome_text = "Welcome to Monthly Report System of Sub Surface Team"
        self.welcome_label = tk.Label(self.root, text=self.welcome_text, font=('Helvetica', 16, 'bold'), background='#0F3057', foreground='yellow')
        self.welcome_label.pack(side=tk.TOP, pady=1.5)
        self.scroll_text()

        # Logo
        self.logo = tk.PhotoImage(file="ongc1.png")
        logo_label = tk.Label(self.root, image=self.logo, bg="#0F3057")
        logo_label.pack(side=tk.TOP, pady=1.5)

        self.frame = ttk.Frame(root, padding="20 20 20 20")
        self.frame.pack(expand=True)
        
        self.create_initial_widgets()
        
        # Initialize user
        self.user = None

        # Employee details
        self.employee_details = pd.DataFrame(columns=["Month", "Year", "Monitoring Group", "Supporting Group","Area Team Details","Miscellaneous Groups", "CPF Number", "Employee Name"])

    def scroll_text(self):
        text = self.welcome_text
        text = text[1:] + text[0]
        self.welcome_label.config(text=text)
        self.welcome_text = text
        self.root.after(150, self.scroll_text)

    def create_initial_widgets(self):
        self.signup_button = ttk.Button(self.frame, text="Sign Up", command=self.create_signup_widgets)
        self.signup_button.grid(row=0, column=0, padx=10, pady=20)

        self.login_button = ttk.Button(self.frame, text="Login", command=self.create_login_widgets)
        self.login_button.grid(row=0, column=1, padx=10, pady=20)

    def create_signup_widgets(self):
        self.clear_frame()
        
        self.email_label = ttk.Label(self.frame, text="Email:")
        self.email_label.grid(row=0, column=0, padx=10, pady=5, sticky="e")
        self.email_entry = ttk.Entry(self.frame)
        self.email_entry.grid(row=0, column=1, padx=10, pady=5)

        self.password_label = ttk.Label(self.frame, text="Password:")
        self.password_label.grid(row=1, column=0, padx=10, pady=5, sticky="e")
        self.password_entry = ttk.Entry(self.frame, show="*")
        self.password_entry.grid(row=1, column=1, padx=10, pady=5)

        self.signup_button = ttk.Button(self.frame, text="Sign Up", command=self.signup)
        self.signup_button.grid(row=2, column=0, columnspan=2, padx=10, pady=20)

        self.back_button = ttk.Button(self.frame, text="Back", command=self.create_initial_widgets)
        self.back_button.grid(row=3, column=0, columnspan=2, padx=10, pady=20)

    def create_login_widgets(self):
        self.clear_frame()
        
        self.email_label = ttk.Label(self.frame, text="Email:")
        self.email_label.grid(row=0, column=0, padx=10, pady=5, sticky="e")
        self.email_entry = ttk.Entry(self.frame)
        self.email_entry.grid(row=0, column=1, padx=10, pady=5)

        self.password_label = ttk.Label(self.frame, text="Password:")
        self.password_label.grid(row=1, column=0, padx=10, pady=5, sticky="e")
        self.password_entry = ttk.Entry(self.frame, show="*")
        self.password_entry.grid(row=1, column=1, padx=10, pady=5)

        self.login_button = ttk.Button(self.frame, text="Login", command=self.login)
        self.login_button.grid(row=2, column=0, columnspan=2, padx=10, pady=20)

        self.back_button = ttk.Button(self.frame, text="Back", command=self.create_initial_widgets)
        self.back_button.grid(row=3, column=0, columnspan=2, padx=10, pady=20)

    def clear_frame(self):
        for widget in self.frame.winfo_children():
            widget.destroy()

    def signup(self):
        email = self.email_entry.get()
        password = self.password_entry.get()

        if not email or not password:
            messagebox.showerror("Error", "All fields are required.")
            return

        # Save user to file
        with open('users.csv', mode='a', newline='') as file:
            writer = csv.writer(file)
            writer.writerow([email, password])

        messagebox.showinfo("Success", "Sign up successful. Please log in.")
        self.create_login_widgets()

    def login(self):
        email = self.email_entry.get()
        password = self.password_entry.get()

        # Check credentials
        if self.check_credentials(email, password):
            self.user = User(email, password)
            self.show_main_frame()
            self.send_welcome_email(email)
        else:
            messagebox.showerror("Login Failed", "Invalid email or password")

    def check_credentials(self, email, password):
        if not os.path.exists('users.csv'):
            return False

        with open('users.csv', mode='r') as file:
            reader = csv.reader(file)
            for row in reader:
                if row[0] == email and row[1] == password:
                    return True
        return False

    def send_welcome_email(self, to_email):
        from_email = "your_email@example.com"
        from_password = "your_email_password"
        subject = "Welcome to SST Monthly Report System"
        body = f"Hello {to_email},\n\nWelcome to the Monthly Report System of the Sub Surface Team (SST). We're glad to have you on board!\n\nBest Regards,\nSST Team"

        msg = MIMEMultipart()
        msg['From'] = from_email
        msg['To'] = to_email
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'plain'))

        try:
            server = smtplib.SMTP('smtp.example.com', 587)  # Replace with your SMTP server and port
            server.starttls()
            server.login(from_email, from_password)
            text = msg.as_string()
            server.sendmail(from_email, to_email, text)
            server.quit()
            messagebox.showinfo("Email Sent", f"Welcome email sent to {to_email}")
        except Exception as e:
            messagebox.showerror("Email Error", f"Failed to send email. Error: {str(e)}")

    def show_main_frame(self):
        self.clear_frame()

        # Manage Employees Frame
        self.manage_frame = tk.Frame(self.root, bd=4, relief=tk.RIDGE, bg="#DAE1E7")
        self.manage_frame.place(x=20, y=100, width=450, height=580)

        m_title = tk.Label(self.manage_frame, text="Manage Employees", bg="#DAE1E7", fg="red", font=("times new roman", 20, "bold"))
        m_title.grid(row=0, columnspan=2, pady=20)

        labels = ["Month", "Year", "Monitoring Group", "Supporting Group","Area Team Details","Miscellaneous Groups", "CPF Number", "Employee Name"]
        self.entries = {}
        for i, label in enumerate(labels):
            lbl = tk.Label(self.manage_frame, text=label, bg="#DAE1E7", fg="black", font=("times new roman", 15, "bold"))
            lbl.grid(row=i+1, column=0, pady=10, padx=10, sticky="w")
            entry = ttk.Entry(self.manage_frame, font=("times new roman", 14, "bold"))
            entry.grid(row=i+1, column=1, pady=10, padx=10, sticky="w")
            self.entries[label] = entry

        self.add_btn = ttk.Button(self.manage_frame, text="Add Employee", command=self.add_employee)
        self.add_btn.grid(row=len(labels) + 1, column=0, pady=10, padx=10)

        self.update_btn = ttk.Button(self.manage_frame, text="Update Employee", command=self.update_employee)
        self.update_btn.grid(row=len(labels) + 1, column=1, pady=10, padx=10)

        self.delete_btn = ttk.Button(self.manage_frame, text="Delete Employee", command=self.delete_employee)
        self.delete_btn.grid(row=len(labels) + 2, column=0, pady=10, padx=10)

        self.clear_btn = ttk.Button(self.manage_frame, text="Clear Fields", command=self.clear_fields)
        self.clear_btn.grid(row=len(labels) + 2, column=1, pady=10, padx=10)

        self.generate_report_btn = ttk.Button(self.manage_frame, text="Generate Report", command=self.generate_report)
        self.generate_report_btn.grid(row=len(labels) + 3, column=0, columnspan=2, pady=10, padx=10)

        # Details Frame
        self.details_frame = tk.Frame(self.root, bd=4, relief=tk.RIDGE, bg="#DAE1E7")
        self.details_frame.place(x=500, y=100, width=680, height=580)

        d_title = tk.Label(self.details_frame, text="Employee Details", bg="#DAE1E7", fg="red", font=("times new roman", 20, "bold"))
        d_title.grid(row=0, columnspan=4, pady=20)

        self.search_label = tk.Label(self.details_frame, text="Search By CPF Number:", bg="#DAE1E7", fg="black", font=("times new roman", 15, "bold"))
        self.search_label.grid(row=1, column=0, pady=10, padx=20, sticky="w")

        self.search_entry = tk.Entry(self.details_frame, width=20, font=("times new roman", 14, "bold"))
        self.search_entry.grid(row=1, column=1, pady=10, padx=20, sticky="w")

        self.search_btn = ttk.Button(self.details_frame, text="Search", command=self.search_employee)
        self.search_btn.grid(row=1, column=2, pady=10, padx=20)

        self.show_all_btn = ttk.Button(self.details_frame, text="Show All", command=self.show_all_employees)
        self.show_all_btn.grid(row=1, column=3, pady=10, padx=20)

        self.table_frame = tk.Frame(self.details_frame, bd=4, relief=tk.RIDGE, bg="#DAE1E7")
        self.table_frame.place(x=10, y=70, width=650, height=500)

        scroll_x = ttk.Scrollbar(self.table_frame, orient=tk.HORIZONTAL)
        scroll_y = ttk.Scrollbar(self.table_frame, orient=tk.VERTICAL)
        self.employee_table = ttk.Treeview(self.table_frame, columns=labels, xscrollcommand=scroll_x.set, yscrollcommand=scroll_y.set)

        scroll_x.pack(side=tk.BOTTOM, fill=tk.X)
        scroll_y.pack(side=tk.RIGHT, fill=tk.Y)
        scroll_x.config(command=self.employee_table.xview)
        scroll_y.config(command=self.employee_table.yview)

        for label in labels:
            self.employee_table.heading(label, text=label)
            self.employee_table.column(label, width=150)

        self.employee_table.pack(fill=tk.BOTH, expand=True)
        self.employee_table.bind("<ButtonRelease-1>", self.get_selected_row)

    def add_employee(self):
        new_data = {label: self.entries[label].get() for label in self.entries}
        self.employee_details = self.employee_details.append(new_data, ignore_index=True)
        self.clear_fields()
        self.show_all_employees()
        messagebox.showinfo("Success", "Employee added successfully")

    def update_employee(self):
        selected_item = self.employee_table.selection()
        if not selected_item:
            messagebox.showerror("Error", "No employee selected")
            return

        employee_index = self.employee_table.index(selected_item[0])
        for label in self.entries:
            self.employee_details.at[employee_index, label] = self.entries[label].get()

        self.clear_fields()
        self.show_all_employees()
        messagebox.showinfo("Success", "Employee updated successfully")

    def delete_employee(self):
        selected_item = self.employee_table.selection()
        if not selected_item:
            messagebox.showerror("Error", "No employee selected")
            return

        employee_index = self.employee_table.index(selected_item[0])
        self.employee_details = self.employee_details.drop(employee_index).reset_index(drop=True)

        self.clear_fields()
        self.show_all_employees()
        messagebox.showinfo("Success", "Employee deleted successfully")

    def search_employee(self):
        cpf_number = self.search_entry.get()
        filtered_data = self.employee_details[self.employee_details["CPF Number"] == cpf_number]

        if filtered_data.empty:
            messagebox.showinfo("Not Found", "No employee found with the given CPF Number")
            return

        self.show_data_in_table(filtered_data)

    def show_all_employees(self):
        self.show_data_in_table(self.employee_details)

    def show_data_in_table(self, data):
        for row in self.employee_table.get_children():
            self.employee_table.delete(row)

        for index, row in data.iterrows():
            self.employee_table.insert("", "end", values=row.tolist())

    def clear_fields(self):
        for entry in self.entries.values():
            entry.delete(0, tk.END)

    def get_selected_row(self, event):
        selected_item = self.employee_table.selection()[0]
        selected_data = self.employee_table.item(selected_item, "values")
        for label, value in zip(self.entries, selected_data):
            self.entries[label].delete(0, tk.END)
            self.entries[label].insert(0, value)

    def clear_email_entry(self):
        self.email_entry.delete(0, tk.END)
        self.password_entry.delete(0, tk.END)

    def generate_report(self):
        self.generate_full_report()

    def generate_full_report(self):
        doc = Document()
        
        # Add front page
        self.add_front_page(doc)

        # Add index page
        self.add_index_page(doc)
        
        # Add employee details in tabular form
        self.add_employee_details_table(doc)
        
        # Add employee details in descriptive form
        self.add_employee_details_descriptive(doc)
        
        # Save document
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if file_path and file_path.endswith('.docx'):
            doc.save(file_path)
            messagebox.showinfo("Report Saved", f"Report saved successfully as {file_path}")
        else:
            messagebox.showerror("Invalid File Type", "Please select a valid .docx file.")

    def add_front_page(self, doc):
        self.add_page_border(doc)
        
        # Create a table to hold the logo and text centered
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Access the single cell in the table
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add the ONGC logo to the cell
        cell.add_paragraph().add_run().add_picture('ongc1.png', width=Inches(2.0))
        doc.add_heading("Monthly Report of Subsurface Team", 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%B %d, %Y')}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_page_break()

    def add_page_border(self, doc):
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

    def add_index_page(self, doc):
        doc.add_heading("Index", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        current_month_year = datetime.now().strftime("%B %Y")
        index_data = [
            {"SI. No.": "1", "PageNo": "1", "TITLE": f"Current Highlights {current_month_year}"},
            {"SI. No.": "2", "PageNo": "2", "TITLE": "Month-wise Physical Targets and Performance"},
            {"SI. No.": "3", "PageNo": "5", "TITLE": f"Work Over Operations {current_month_year}"},
            {"SI. No.": "4", "PageNo": "6", "TITLE": "Reservoir Field Operations"},
            {"SI. No.": "5", "PageNo": "6", "TITLE": "Pressure and Temperature Measurements"},
            {"SI. No.": "6", "PageNo": "7", "TITLE": "Well-wise Pressure & Temperature Data"},
            {"SI. No.": "7", "PageNo": "11", "TITLE": "Gas Production"},
            {"SI. No.": "7.1", "PageNo": "13", "TITLE": "Sand wise/ Well wise Gas Production"},
            {"SI. No.": "8", "PageNo": "18", "TITLE": f"Pay Zone Wise Status of Wells as on {current_month_year}"},
            {"SI. No.": "9", "PageNo": "20", "TITLE": "Production Test data"},
            {"SI. No.": "9.1", "PageNo": "25", "TITLE": "Latest Production Test Data of Flowing Wells"},
            {"SI. No.": "10", "PageNo": "27", "TITLE": "Pressure Production data (Initial testing) of unconnected wells"},
            {"SI. No.": "11", "PageNo": "28", "TITLE": "Development Locations available for drilling"},
            {"SI. No.": "12", "PageNo": "29", "TITLE": f"Status of Petroleum Mining Lease {current_month_year}"}
        ]

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'SI. No.'
        hdr_cells[1].text = 'PageNo'
        hdr_cells[2].text = 'TITLE'

        for item in index_data:
            row_cells = table.add_row().cells
            row_cells[0].text = item['SI. No.']
            row_cells[1].text = item['PageNo']
            row_cells[2].text = item['TITLE']

        doc.add_page_break()

    def add_employee_details(self, doc):
        doc.add_heading('Employee Details', level=1)
        table = doc.add_table(rows=1, cols=len(self.employee_details.columns))
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        for i, column in enumerate(self.employee_details.columns):
            hdr_cells[i].text = column

        for _, row in self.employee_details.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

        doc.add_page_break()

    def add_excel_sheet_data(self, doc):
        doc.add_heading("Monthly Report Of SST", 0)
        current_date = datetime.now().strftime("%Y-%m-%d")
        doc.add_paragraph(f"Date: {current_date}", style='Intense Quote')

        for _, row in self.employee_details.iterrows():
            p = doc.add_paragraph()
            for label, value in row.items():
                p.add_run(f"{label}: ").bold = True
                p.add_run(f"{value}\n")


        doc.add_page_break()

        # Add Employee Details section in tabular form
        employees = self.employee_details.to_dict(orient='records')
        table = doc.add_table(rows=1, cols=len(self.employee_details.columns))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(self.employee_details.columns):
            hdr_cells[i].text = col
        for emp in employees:
            row_cells = table.add_row().cells
            for i, col in enumerate(self.employee_details.columns):
                row_cells[i].text = str(emp[col])

        # Add Employee Details section in descriptive form
        doc.add_heading("Employee Details", level=2)
        for _, row in self.employee_details.iterrows():
            for col in self.employee_details.columns:
                doc.add_paragraph(f"{col}: {row[col]}")
            doc.add_paragraph()  # Add an empty paragraph for spacing


def add_excel_sheet_data(self, doc):
        # Function to extract text from an image using OCR
        def extract_text_from_image(image_path):
            image = image.open(image_path)
            text = pytesseract.image_to_string(image)
            return text

        # Example of extracting data from the provided images
        image_paths = ['Excel_sheet1_screenshot.jpg', 'Excel_sheet2_screenshot.jpg']

        for image_path in image_paths:
            extracted_text = extract_text_from_image(image_path)
            # Assuming the text is tabular data separated by tabs or newlines
            rows = extracted_text.split('\n')
            if rows:
                table = doc.add_table(rows=1, cols=len(rows[0].split('\t')))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(rows[0].split('\t')):
                    hdr_cells[i].text = col
                for row in rows[1:]:
                    row_cells = table.add_row().cells
                    for i, col in enumerate(row.split('\t')):
                        row_cells[i].text = col
                doc.add_paragraph()  # Add spacing after the table

        doc.add_page_break()

def add_employee_details_table(self, doc):
        doc.add_heading("Employee Details", level=1)
        if self.employee_details.empty:
            doc.add_paragraph("No employee details available.")
        else:
            table = doc.add_table(rows=1, cols=len(self.employee_details.columns))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, column in enumerate(self.employee_details.columns):
                hdr_cells[i].text = column

            for _, row in self.employee_details.iterrows():
                row_cells = table.add_row().cells
                for i, cell in enumerate(row):
                    row_cells[i].text = str(cell)
        
        doc.add_page_break()

def add_employee_details_descriptive(self, doc):
        doc.add_heading("Employee Details Descriptive", level=1)
        if self.employee_details.empty:
            doc.add_paragraph("No employee details available.")
        else:
            for _, row in self.employee_details.iterrows():
                for column, value in row.items():
                    doc.add_heading(column, level=2)
                    doc.add_paragraph(str(value))
                doc.add_paragraph("\n")
        
        doc.add_page_break()

        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if file_path and file_path.endswith('.docx'):
            doc.save(file_path)
            messagebox.showinfo("Report Saved", f"Report saved successfully as {file_path}")
        else:
            messagebox.showerror("Invalid File Type", "Please select a valid .docx file.")

def main():
    root = tk.Tk()
    app = MonthlyReportSystem(root)
    root.mainloop()

if __name__ == "__main__":
    main()

